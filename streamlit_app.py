import streamlit as st
import io
import sys
import os
from pypdf import PdfReader
from docx import Document
import pandas as pd
import numpy as np
import datetime
from fpdf import FPDF

# 确保能找到 ai_core 模块
project_root = os.path.dirname(os.path.abspath(__file__))
if project_root not in sys.path:
    sys.path.append(project_root)

from ai_core.rag_engine import analyze_document_with_ai

def extract_text_from_file(file_content: bytes, filename: str) -> str:
    """从不同格式的文件中提取文本"""
    text = ""
    file_extension = filename.split('.')[-1].lower()
    try:
        if file_extension == 'pdf':
            reader = PdfReader(io.BytesIO(file_content))
            for page in reader.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
        elif file_extension in ['docx', 'doc']:
            doc = Document(io.BytesIO(file_content))
            for para in doc.paragraphs:
                text += para.text + "\n"
        else:
            text = file_content.decode('utf-8', errors='ignore')
    except Exception as e:
        st.error(f"文件解析失败: {str(e)}")
        return ""
    
    return text.strip()

def generate_docx_report(report: dict, filename: str) -> bytes:
    doc = Document()
    doc.add_heading(f"【AI特检】{filename} - 审查诊断报告", 0)
    
    risks = report.get("合规性风险", [])
    if risks:
        doc.add_heading("🔴 合规风险 (必须整改)", level=1)
        for i, r in enumerate(risks, 1):
            doc.add_heading(f"风险 {i}: {r.get('描述', '未知')}", level=2)
            p = doc.add_paragraph()
            p.add_run("⚡ AI 修复建议：").bold = True
            p.add_run(r.get('建议', '无'))
            
    logics = report.get("逻辑错误", [])
    if logics:
        doc.add_heading("🟡 逻辑异常 (自相矛盾)", level=1)
        for i, l in enumerate(logics, 1):
            doc.add_heading(f"异常 {i}: {l.get('描述', '未知')}", level=2)
            p = doc.add_paragraph()
            p.add_run("⚡ AI 修复建议：").bold = True
            p.add_run(l.get('建议', '无'))
            
    infos = report.get("核心信息", [])
    if infos:
        doc.add_heading("🔵 核心提取信息", level=1)
        for i, info in enumerate(infos, 1):
            doc.add_paragraph(f"{info.get('描述', '未知')}：{info.get('建议', '')}", style='List Bullet')
            
    p = doc.add_paragraph("\n")
    run = p.add_run("华招国际内部参阅专用")
    run.bold = True
    
    bio = io.BytesIO()
    doc.save(bio)
    return bio.getvalue()

def generate_pdf_reminder(project_name: str, manager: str) -> bytes:
    pdf = FPDF()
    pdf.add_page()
    
    font_path = "/Library/Fonts/Arial Unicode.ttf"
    if os.path.exists(font_path):
        pdf.add_font("ArialUnicode", "", font_path)
        pdf.set_font("ArialUnicode", "", 14)
    else:
        pdf.set_font("Arial", "", 14)
        
    pdf.cell(200, 10, text=f"催办函 - {project_name}", align="C")
    pdf.ln(15)
    
    body_text = (
        f"尊敬的 {manager} 负责人：\n\n"
        f"您好！\n\n"
        f"系统检测到【{project_name}】项目的电子化进度低于 50%，"
        f"且距离开标日期已不足 3 天。为确保项目顺利推进，请您立即跟进处理，"
        f"加快电子化流程，并在系统内更新最新进度。\n\n"
        f"特此提醒！\n\n"
        f"自动预警系统\n"
        f"生成时间: {datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S')}"
    )
    pdf.multi_cell(0, 10, text=body_text)
    
    return bytes(pdf.output())

st.set_page_config(page_title="系统 | AI 辅助招标及实施预警", layout="wide")

CUSTOM_CSS = """
<style>
    /* 移除容易导致重叠的 padding 和溢出隐藏设置，保留安全的字体和内部排版体系 */
    @import url('https://fonts.googleapis.com/css2?family=Inter:wght@400;500;600&family=Manrope:wght@600;700;800&display=swap');
    
    html, body, [class*="css"] {
        font-family: 'Inter', sans-serif;
    }
    h1, h2, h3, h4, h5, h6 {
        font-family: 'Manrope', sans-serif !important;
        color: #00174b !important;
    }
    
    /* 主背景色轻微区分 */
    .stApp {
        background-color: #f8fafc;
    }
    
    /* AI 修复建议专属高亮框 (安全的 HTML 隔离区) */
    .ai-suggestion-box {
        background-color: #f1f5f9;
        border-radius: 6px;
        padding: 12px 16px;
        margin-top: 8px;
        margin-bottom: 8px;
        border-left: 4px solid #0052d9;
        color: #1e293b;
    }
    .ai-suggestion-box strong {
        color: #003da6;
    }
    
    /* 结果分析卡片 (Expander) 无损轻量美化 */
    div[data-testid="stExpander"] {
        border-left: 4px solid #003da6 !important;
        border-radius: 8px !important;
        box-shadow: 0 2px 12px rgba(0, 61, 166, 0.05) !important;
    }
</style>
"""
st.markdown(CUSTOM_CSS, unsafe_allow_html=True)

st.sidebar.markdown("## 🧰 华招国际政企业务中台")
nav_choice = st.sidebar.radio(
    "请选择场景模块：",
    ["📄 招标文件智能审查 (文本级)", "⚠️ 交易系统操作预警 (实操级)", "📊 AI 智能协作平台 (管理级)"]
)

st.sidebar.divider()
st.sidebar.header("⚙️ 引擎流转状态 (云端)")
st.sidebar.info("✓ 建立通道")
st.sidebar.info("✓ RAG 预热完毕")
st.sidebar.info("✓ LLM 推理就绪")

if nav_choice == "📄 招标文件智能审查 (文本级)":
    st.markdown("<h1 style='color: #003da6;'>🤖 智能审查大脑 - 招标文件自动诊断</h1>", unsafe_allow_html=True)
    st.markdown("<p style='color: #434654; font-size: 1.1rem;'>通过大语言模型深度理解和 RAG 技术，一键识别标书中的合规性风险与前后逻辑冲突。</p>", unsafe_allow_html=True)
    
    # --- 主区 ---
    uploaded_file = st.file_uploader("请上传待审查的招标文件 (支持 PDF / Word / TXT 格式)", type=["pdf", "docx", "txt", "doc"])

    if uploaded_file is not None:
        st.success(f"文件加载就绪：{uploaded_file.name}")
        
        if st.button("🚀 开始 AI 智能审查", type="primary", use_container_width=True):
            with st.status(">> 正在启动 AI 审查引擎...", expanded=True) as status:
                st.write(">> 阶段 1：正在提取分析文档文本...")
                
                # 直接在本地调用提取
                extracted_text = extract_text_from_file(uploaded_file.getvalue(), uploaded_file.name)
                
                if not extracted_text:
                    status.update(label="审查失败，无法提取文件文本", state="error", expanded=True)
                else:
                    st.write(">> 阶段 2：大模型正在进行语义理解与合规推断 (这可能会耗时十秒以上，请耐心等待)...")
                    
                    # 直接调用大模型核心逻辑
                    review_result = analyze_document_with_ai(extracted_text)
                    
                    # 判断是否有报错返回
                    if review_result.get("status") == "error":
                        status.update(label="审查失败，AI 模型调用异常", state="error", expanded=True)
                        st.error(review_result.get("message"))
                        st.write(review_result.get("error_detail"))
                    else:
                        status.update(label="审查完毕，成功生成结构化报告。", state="complete", expanded=False)
                        
                        report = review_result
                        
                        st.divider()
                        st.subheader(f"📊 【{uploaded_file.name}】自动审查诊断报告", divider="red")
                        
                        tab1, tab2, tab3 = st.tabs(["🔴 合规风险", "🟡 逻辑异常", "🔵 核心提取信息"])
                        
                        with tab1:
                            risks = report.get("合规性风险", [])
                            if risks:
                                st.error(f"AI 识别出 **{len(risks)}** 个合规性风险点（**触及红线，必须整改**）")
                                for idx, risk in enumerate(risks):
                                    with st.expander(f"⚠️ 风险 {idx+1}: {risk.get('描述', '未知')}", expanded=True):
                                        st.markdown(f"<div class='ai-suggestion-box'><strong>⚡ AI 修复建议：</strong><br>{risk.get('建议', '无')}</div>", unsafe_allow_html=True)
                            else:
                                st.success("✅ 未发现明显的合规性风险。")
        
                        with tab2:
                            logics = report.get("逻辑错误", [])
                            if logics:
                                st.warning(f"AI 识别出 **{len(logics)}** 处前后逻辑或数据自相矛盾")
                                for idx, logic in enumerate(logics):
                                    with st.expander(f"📌 异常 {idx+1}: {logic.get('描述', '未知')}", expanded=True):
                                        st.markdown(f"<div class='ai-suggestion-box'><strong>⚡ AI 修复建议：</strong><br>{logic.get('建议', '无')}</div>", unsafe_allow_html=True)
                            else:
                                st.success("✅ 未发现明显的逻辑矛盾。")
                                        
                        with tab3:
                            infos = report.get("核心信息", [])
                            if infos:
                                st.info(f"AI 提取了以下关键项目要素：")
                                for idx, info in enumerate(infos):
                                    st.markdown(f"- **{info.get('描述', '未知')}**：{info.get('建议', '')}")
                            else:
                                st.info("暂未提取到具体的要素记录。")
                                
                        st.divider()
                        st.markdown("### 💾 导出审查中心报告")
                        docx_data = generate_docx_report(report, uploaded_file.name)
                        st.download_button(
                            label="📥 下载 Word (.docx) 格式报告文件",
                            data=docx_data,
                            file_name=f"【AI特检报告】_{uploaded_file.name}.docx",
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                            use_container_width=True
                        )
                        
                        st.divider()
                        with st.expander("开发者模式：查看服务端返回的完整 JSON 原始结构"):
                            st.json(report)

elif nav_choice == "⚠️ 交易系统操作预警 (实操级)":
    from ai_core.rag_engine import operation_warning_agent
    
    st.title("⚠️ 交易平台与工具排障向导")
    st.markdown("专治【甘肃省公共资源交易电子服务系统】 / 【离线固化客户端】 / 【CA 互认互签】等各种疑难杂症与导致废标的诡异拦截。")
    st.info("💡 提示：本模块直连后端 2026 年最新排障数据库，请准确描述您看到的网页红色报错或弹框文字。")
    
    user_issue = st.text_area("✍️ 请在此描述您当前急需解决的系统报错、操作卡顿或疑惑（例如：'交了钱进不去开标大厅' 或 'HASH值生成失败'）", height=150)
    
    if st.button("🚨 召唤系统运维专家查阅", type="primary", use_container_width=True):
        if user_issue.strip():
            with st.spinner("🕵️‍♂️ 专家正在急速检索关联异常库档案，请稍作等待..."):
                diag = operation_warning_agent(user_issue)
                
            st.success("✅ 诊断出具完毕！请核对自己所处的处境：")
            
            col1, col2 = st.columns(2)
            with col1:
                st.markdown("### 🔍 根本原因确诊 (Diagnosis)")
                st.info(diag.get("Diagnosis", "未获取到信息"))
                
                st.markdown("### 💥 严重后果预警 (Warning)")
                st.error(diag.get("Warning", "未获取到信息"))
            
            with col2:
                st.markdown("### 💊 紧急抢救指南 (Action Plan)")
                st.success(diag.get("ActionPlan", "未获取到信息"))
        else:
            st.warning("⚠️ 请务必键入您的问题。")

elif nav_choice == "📊 AI 智能协作平台 (管理级)":
    st.title("🚀 招标业务线上化转型监控")
    st.subheader("5人核心团队任务执行状态")

    # 模拟数据：各项目进度
    chart_data = pd.DataFrame(
        np.random.randn(20, 3),
        columns=['电子标书制作', '在线开标模拟', '系统集成测试']
    )

    # 1. 核心指标统计
    col1, col2, col3 = st.columns(3)
    col1.metric("当月在线招标项目", "42 个", "12%")
    col2.metric("核心团队负荷", "85%", "-5%")
    col3.metric("省财政厅接口通畅度", "稳定", "99.9%")

    # 2. 将控制组件放入 Sidebar
    st.sidebar.divider()
    st.sidebar.header("管理筛选")
    member = st.sidebar.selectbox("查看技术人员：", ["张三 (前端)", "李四 (后端)", "王五 (AI预审)", "赵六 (测试)", "孙七 (运维)"])
    st.sidebar.write(f"当前正在查看 **{member}** 的工作流。")

    # 3. 进度可视化
    st.line_chart(chart_data)

    # 4. 干系人满意度分析
    st.divider()
    st.subheader("🤝 干系人（业主/专家）满意度分析")
    satisfaction = st.slider("当前项目平均满意度", 0, 100, 85)
    if satisfaction < 60:
        st.error("警告：满意度低于及格线，请启动『干系人参与计划』！")
    else:
        st.success("状态良好：继续保持沟通。")

    # 5. 项目预警系统
    st.divider()
    st.subheader("🚨 智能项目预警系统")
    
    csv_path = os.path.join(project_root, "projects.csv")
    if os.path.exists(csv_path):
        df = pd.read_csv(csv_path)
        
        # 数据处理
        df['开标日期'] = pd.to_datetime(df['开标日期'])
        df['距离天数'] = (df['开标日期'] - pd.Timestamp('today')).dt.days
        
        # 判断警告条件
        df['预警'] = (df['电子化进度'] < 50) & (df['距离天数'] <= 3)

        # 样式高亮函数
        def highlight_warning(row):
            color = 'background-color: #ffcccc' if row['预警'] else ''
            return [color] * len(row)

        # 展示带有颜色的表格
        st.markdown("##### 当前所有项目电子化进度监控：")
        st.dataframe(df.style.apply(highlight_warning, axis=1), use_container_width=True)
        
        # 筛选出需要警告的项目生成催办函
        warning_projects = df[df['预警']]
        if not warning_projects.empty:
            st.warning(f"发现 {len(warning_projects)} 个项目进度滞后且临近开标，建议立即生成发送催办函！")
            
            for idx, row in warning_projects.iterrows():
                proj_name = row['项目名称']
                mgr = row['负责人']
                
                pdf_bytes = generate_pdf_reminder(proj_name, mgr)
                
                # 使用下载按钮提供下载
                st.download_button(
                    label=f"📥 下载《{proj_name}》PDF催办函",
                    data=pdf_bytes,
                    file_name=f"{proj_name}_催办函.pdf",
                    mime="application/pdf",
                    key=f"dl_btn_{idx}"
                )
        else:
            st.success("太棒了！目前所有项目进度均在计划范围内，或距离开标日期时间充足。")
    else:
        st.info("未找到 `projects.csv` 数据文件，请确保将其放置在同级目录下。")

